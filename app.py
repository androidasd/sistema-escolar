import streamlit as st
import pandas as pd
from docx import Document
from github import Github, Auth  # <--- Adicionado Auth aqui
import io
import time
import json
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from streamlit_option_menu import option_menu

# --- CONFIGURAÇÃO DA PÁGINA (PRIMEIRA LINHA OBRIGATÓRIA) ---
ST_COR_PADRAO = "#00A8C6"
ST_TITULO_PADRAO = "SISTEMA ESCOLAR"

st.set_page_config(page_title=ST_TITULO_PADRAO, page_icon="🎓", layout="wide")

# --- CSS NUCLEAR (A SOLUÇÃO DEFINITIVA) ---
st.markdown("""
    <style>
        /* Importa fonte */
        @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;700&display=swap');
        html, body, [class*="css"] { font-family: 'Roboto', sans-serif; }

        /* --- ZONA DE SUPRESSÃO TOTAL DO STREAMLIT --- */
        
        /* 1. Mata o cabeçalho nativo completamente e remove o espaço dele */
        header {
            visibility: hidden !important;
            display: none !important;
            height: 0px !important;
            opacity: 0 !important;
            pointer-events: none !important;
        }
        
        /* 2. Remove especificamente a barra de ferramentas e botões de deploy */
        [data-testid="stToolbar"], 
        [data-testid="stAppDeployButton"], 
        [data-testid="stDecoration"],
        [data-testid="stHeader"] {
            visibility: hidden !important;
            display: none !important;
            height: 0px !important;
        }

        /* 3. Remove rodapé e menu */
        footer, #MainMenu {
            display: none !important;
            visibility: hidden !important;
        }
        
        /* 4. Remove o badge "Viewer" que as vezes aparece embaixo */
        .viewerBadge_container__1QSob {
            display: none !important;
        }

        /* 5. AJUSTE CRÍTICO: Sobe o conteúdo para tapar o buraco do header */
        .block-container {
            padding-top: 0rem !important;
            margin-top: -4rem !important; /* Puxa tudo para cima agressivamente */
            padding-bottom: 2rem !important;
        }
        
        /* --- ESTILOS DO SEU SISTEMA --- */
        :root {
            --primary: #00A8C6;
            --card-bg: #ffffff;
        }

        /* Login Container Compacto */
        .login-container {
            background: white;
            padding: 30px;
            border-radius: 15px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.08);
            border: 1px solid #eee;
            margin-top: 60px; /* Dá um respiro agora que subimos o container */
        }

        /* Botões */
        div.stButton > button:first-child {
            background-color: var(--primary);
            color: white;
            border: none;
            border-radius: 8px;
            font-weight: 600;
        }
        div.stButton > button:first-child:hover { opacity: 0.9; }

        /* Cards de Métricas */
        div[data-testid="metric-container"] {
            background-color: var(--card-bg);
            border-radius: 10px;
            padding: 15px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            border-left: 5px solid var(--primary);
            text-align: center;
        }
        
        /* Popup do Perfil */
        .profile-popup-box {
            background-color: white;
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            box-shadow: 0 4px 12px rgba(0,0,0,0.15);
            color: #333;
            margin-top: 5px;
            font-size: 14px;
        }
    </style>
""", unsafe_allow_html=True)

# --- FUNÇÕES ---
def hash_senha(senha):
    return hashlib.sha256(str.encode(senha)).hexdigest()

def enviar_email_boas_vindas(destinatario, nome_usuario):
    try:
        remetente = st.secrets["EMAIL_USER"]
        senha_app = st.secrets["EMAIL_PASSWORD"].replace(" ", "").strip()
        remetente = remetente.strip()
        msg = MIMEMultipart()
        msg['From'] = remetente
        msg['To'] = destinatario
        msg['Subject'] = "Acesso Solicitado - Sistema Escolar"
        texto = f"""Olá, {nome_usuario}!\nRecebemos sua solicitação de cadastro.\nLogin: {destinatario}\nSituação: PENDENTE."""
        msg.attach(MIMEText(texto, 'plain'))
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(remetente, senha_app)
        server.sendmail(remetente, destinatario, msg.as_string())
        server.quit()
        return True, "Enviado"
    except Exception as e: return False, str(e)

# --- CONEXÃO GITHUB (CORRIGIDA PARA NOVO PADRÃO) ---
try:
    TOKEN = st.secrets["GITHUB_TOKEN"]
    # Correção do DeprecationWarning: Usando Auth.Token
    auth = Auth.Token(TOKEN)
    g = Github(auth=auth)
    
    user = g.get_user()
    repo_ref = None
    for repo in user.get_repos():
        if "sistema" in repo.name.lower() or "escolar" in repo.name.lower() or "emeif" in repo.name.lower():
            repo_ref = repo; break
    if not repo_ref: 
        repos = list(user.get_repos()); 
        if repos: repo_ref = repos[0]
    if not repo_ref: st.error("Erro Crítico: Repositório não encontrado."); st.stop()
except Exception as e: st.error(f"Erro de conexão: {e}"); st.stop()

# --- ARQUIVOS ---
ARQ_PASSIVOS = 'EMEF PA-RESSACA.docx'
ARQ_CONCLUINTES = 'CONCLUINTES- PA-RESSACA.docx'
ARQ_USERS = 'users.json'
ARQ_CONFIG = 'config.json'

# --- MANIPULAÇÃO DADOS ---
def carregar_json(arquivo):
    try:
        content = repo_ref.get_contents(arquivo)
        return json.loads(content.decoded_content.decode()), content.sha
    except: return {}, None

def salvar_json(arquivo, dados, sha, mensagem):
    try:
        dados_str = json.dumps(dados, indent=4)
        if sha: repo_ref.update_file(arquivo, mensagem, dados_str, sha)
        else: repo_ref.create_file(arquivo, mensagem, dados_str)
        return True
    except: return False

@st.cache_data(ttl=60)
def carregar_dados_word():
    lista = []
    def processar(nome_arq, categoria):
        local = []
        try:
            c = repo_ref.get_contents(nome_arq)
            doc = Document(io.BytesIO(c.decoded_content))
            for tabela in doc.tables:
                for linha in tabela.rows:
                    if len(linha.cells) >= 2:
                        num = linha.cells[0].text.strip()
                        nome = linha.cells[1].text.strip().upper()
                        obs = linha.cells[2].text.strip() if len(linha.cells) > 2 else ""
                        if len(nome) > 3 and "NOME" not in nome:
                            local.append({"Numero": num, "Nome": nome, "Categoria": categoria, "Obs": obs})
            return local
        except: return []
    return processar(ARQ_PASSIVOS, "Passivo") + processar(ARQ_CONCLUINTES, "Concluinte")

def salvar_aluno_word(arquivo_nome, numero, nome, obs):
    try:
        c = repo_ref.get_contents(arquivo_nome)
        doc = Document(io.BytesIO(c.decoded_content))
        if len(doc.tables) > 0:
            tab = doc.tables[0]
            row = tab.add_row()
            row.cells[0].text = numero
            row.cells[1].text = nome.upper()
            if len(row.cells) > 2: row.cells[2].text = obs
            buffer = io.BytesIO()
            doc.save(buffer)
            repo_ref.update_file(arquivo_nome, f"Add Aluno: {nome}", buffer.getvalue(), c.sha)
            return True
    except: return False

config_data, config_sha = carregar_json(ARQ_CONFIG)
COR_TEMA = config_data.get("theme_color", ST_COR_PADRAO)
NOME_ESCOLA = config_data.get("school_name", ST_TITULO_PADRAO)
LOGO_URL = config_data.get("logo_url", "https://cdn-icons-png.flaticon.com/512/3135/3135715.png")

# --- LÓGICA DE SESSÃO ---
if 'user_info' not in st.session_state: st.session_state['user_info'] = None

# ==============================================================================
# TELA DE LOGIN (SUPER COMPACTA)
# ==============================================================================
if not st.session_state['user_info']:
    # AJUSTE: [5, 3, 5] -> A coluna do meio (3) é a única com conteúdo
    col_e, col_c, col_d = st.columns([5, 3, 5])
    
    with col_c:
        with st.container():
            st.markdown(f"""
            <div class="login-container" style="text-align:center;">
                <img src="{LOGO_URL}" width="70" style="margin-bottom:10px;">
                <h3 style="color:{COR_TEMA}; margin:0; font-weight:700;">{NOME_ESCOLA}</h3>
                <p style="color:gray; font-size:12px;">Gestão Acadêmica</p>
                <hr style="opacity:0.2; margin: 15px 0;">
            </div>
            """, unsafe_allow_html=True)
            
            tab1, tab2 = st.tabs(["ENTRAR", "CADASTRAR"])
            
            with tab1:
                with st.form("login_frm"):
                    email = st.text_input("E-mail")
                    senha = st.text_input("Senha", type="password")
                    if st.form_submit_button("ACESSAR", use_container_width=True):
                        try: s_adm = st.secrets["SENHA_SISTEMA"]
                        except: s_adm = "admin"
                        # ADMIN: admin@gmail.com
                        if email.lower() == "admin@gmail.com" and senha == s_adm:
                            st.session_state['user_info'] = {"username": "Admin", "name": "Administrador Principal", "role": "admin", "email": "admin@gmail.com", "unit": "DIRETORIA"}
                            st.rerun()
                        # USUARIOS NORMAIS
                        db, _ = carregar_json(ARQ_USERS)
                        u = next((x for x in db.get("users", []) if x.get('email', '').lower() == email.lower() and x['password'] == hash_senha(senha)), None)
                        if u:
                            if u.get('status') == 'active': st.session_state['user_info'] = u; st.rerun()
                            else: st.warning("Cadastro em análise.")
                        else: st.error("Dados inválidos.")
            with tab2:
                with st.form("reg_frm"):
                    n = st.text_input("Nome"); e = st.text_input("E-mail"); s = st.text_input("Senha", type="password")
                    if st.form_submit_button("CRIAR CONTA", use_container_width=True):
                        if "@" not in e: st.error("E-mail inválido")
                        else:
                            db, sha = carregar_json(ARQ_USERS)
                            lst = db.get("users", [])
                            if any(x.get('email') == e for x in lst): st.error("E-mail já existe.")
                            else:
                                with st.spinner("Criando..."):
                                    lst.append({"username": e.split("@")[0], "password": hash_senha(s), "name": n, "email": e, "role": "user", "status": "pending", "unit": "Geral"})
                                    if not db: db = {"users": []}
                                    db['users'] = lst
                                    salvar_json(ARQ_USERS, db, sha, f"Reg {e}")
                                    env, erro = enviar_email_boas_vindas(e, n)
                                    if env: st.success("Solicitação enviada!")
                                    else: st.warning(f"Salvo (Sem email)")
    st.stop()

# ==============================================================================
# ÁREA LOGADA
# ==============================================================================
user = st.session_state['user_info']

# --- HEADER ---
with st.container():
    c_logo, c_user = st.columns([2, 3])
    with c_logo:
        st.markdown(f"""
        <div style="display:flex; align-items:center; gap:12px;">
            <img src="{LOGO_URL}" width="40">
            <div><h4 style="margin:0; color:{COR_TEMA}; font-weight:800;">{NOME_ESCOLA}</h4></div>
        </div>
        """, unsafe_allow_html=True)
    with c_user:
        c_info, c_logout = st.columns([4, 1])
        with c_info:
            with st.expander(f"👤 {user['name'].split()[0]}", expanded=False):
                st.markdown(f"""<div class="profile-popup-box"><strong>{user['name']}</strong><br><small>{user.get('email')}</small><br><span style="color:{COR_TEMA}; font-weight:bold;">{user['role'].upper()}</span></div>""", unsafe_allow_html=True)
        with c_logout:
            if st.button("SAIR"): st.session_state['user_info'] = None; st.rerun()

st.divider()

# --- MENU HORIZONTAL ---
opts = ["Dashboard", "Pesquisar", "Cadastrar Aluno"]
icons = ["house", "search", "person-plus"]
if user['role'] == 'admin': opts.append("Administração"); icons.append("gear")

selected = option_menu(
    menu_title=None, options=opts, icons=icons, default_index=0, orientation="horizontal",
    styles={
        "container": {"padding": "0!important", "background-color": "#ffffff", "border-radius": "5px"},
        "icon": {"color": COR_TEMA, "font-size": "15px"},
        "nav-link": {"font-size": "14px", "text-align": "center", "margin": "0px", "--hover-color": "#f0f2f6"},
        "nav-link-selected": {"background-color": COR_TEMA, "color": "white"},
    }
)
st.write("")

# --- PÁGINAS ---
if selected in ["Dashboard", "Pesquisar"]: df = pd.DataFrame(carregar_dados_word())

if selected == "Dashboard":
    st.subheader("📊 Visão Geral")
    if not df.empty:
        col1, col2, col3 = st.columns(3)
        col1.metric("Total", len(df)); col2.metric("Concluintes", len(df[df['Categoria']=="Concluinte"])); col3.metric("Passivos", len(df[df['Categoria']=="Passivo"]))
        st.write(""); st.markdown("##### 📌 Últimas Atualizações")
        st.dataframe(df.tail(8), use_container_width=True, hide_index=True)
    else: st.info("Sem dados.")

elif selected == "Pesquisar":
    st.subheader("🔍 Buscar Aluno")
    busca = st.text_input("Digite o nome...", placeholder="Ex: Maria da Silva")
    if busca and not df.empty:
        res = df[df['Nome'].str.contains(busca.upper(), na=False)]
        if not res.empty: st.success(f"{len(res)} encontrados."); st.dataframe(res, use_container_width=True, hide_index=True)
        else: st.warning("Não encontrado.")

elif selected == "Cadastrar Aluno":
    st.subheader("📝 Nova Matrícula")
    with st.container():
        with st.form("novo_aluno_form"):
            c1, c2 = st.columns([1, 4])
            num = c1.text_input("Nº Chamada", placeholder="000")
            nome = c2.text_input("Nome Completo")
            c3, c4 = st.columns(2)
            tipo = c3.radio("Situação", ["Passivos", "Concluintes"], horizontal=True)
            obs = c4.text_input("Observação")
            st.write("")
            if st.form_submit_button("💾 SALVAR ALUNO", use_container_width=True):
                arq = ARQ_PASSIVOS if tipo == "Passivos" else ARQ_CONCLUINTES
                if not num: num = "S/N"
                if salvar_aluno_word(arq, num, nome, obs): st.toast(f"Salvo!", icon="✅"); time.sleep(1); st.cache_data.clear(); st.rerun()
                else: st.error("Erro ao salvar.")

elif selected == "Administração":
    st.subheader("⚙️ Configurações")
    tab_users, tab_pass, tab_config = st.tabs(["Usuários", "Senhas", "Sistema"])
    with tab_users:
        db, sha = carregar_json(ARQ_USERS)
        if db.get("users"):
            users_df = pd.DataFrame(db['users'])
            cols = ["name", "email", "status", "role"]
            show_df = users_df[[c for c in cols if c in users_df.columns]]
            edited = st.data_editor(show_df, key="editor_users", use_container_width=True,
                column_config={"status": st.column_config.SelectboxColumn("Acesso", options=["active", "pending", "disabled"]),
                               "role": st.column_config.SelectboxColumn("Nível", options=["user", "admin"])})
            if st.button("Salvar Acessos"):
                novos = edited.to_dict('records')
                lista_final = []
                for n in novos:
                    orig = next((u for u in db['users'] if u['email'] == n['email']), None)
                    if orig: orig.update(n); lista_final.append(orig)
                    else: lista_final.append(n)
                db['users'] = lista_final
                salvar_json(ARQ_USERS, db, sha, "Update Users"); st.success("Atualizado!"); time.sleep(1); st.rerun()
    with tab_pass:
        st.write("Trocar senha:")
        db, sha = carregar_json(ARQ_USERS)
        lst = db.get("users", [])
        sel_user = st.selectbox("Usuário:", [u['email'] for u in lst])
        if sel_user:
            p1 = st.text_input("Nova Senha", type="password"); p2 = st.text_input("Repetir Senha", type="password")
            if st.button("Confirmar Troca"):
                if p1 == p2:
                    for u in lst:
                        if u['email'] == sel_user: u['password'] = hash_senha(p1)
                    db['users'] = lst
                    salvar_json(ARQ_USERS, db, sha, "Update pass"); st.success("Senha atualizada!")
                else: st.error("Senhas não conferem.")
    with tab_config:
        st.write("Aparência:")
        with st.form("conf_geral"):
            cn = st.text_input("Nome", NOME_ESCOLA); cc = st.color_picker("Cor", COR_TEMA); cl = st.text_input("Logo", LOGO_URL)
            if st.form_submit_button("Salvar"):
                _, s_c = carregar_json(ARQ_CONFIG)
                salvar_json(ARQ_CONFIG, {"school_name": cn, "theme_color": cc, "logo_url": cl}, s_c, "Upd Config"); st.rerun()
